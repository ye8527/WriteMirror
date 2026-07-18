"""Embed architecture and organizer-facing evidence into the Japanese DOCX files."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PLAN = next(DOCS.glob("*Vibe_Coding*.docx"))
GUIDE = next(DOCS.glob("*利用・技術説明書*.docx"))
DIAGRAM = DOCS / "assets" / "WriteMirror_システムアーキテクチャ_ja.png"
CAPTION = "図1　WriteMirror 0.6.0の入力、端末内処理、出力、学習経路、通信境界"


def find(document: Document, fragment: str, *, last: bool = False) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if fragment in paragraph.text]
    if not matches:
        raise ValueError(f"Paragraph not found: {fragment}")
    return matches[-1] if last else matches[0]


def insert_after(paragraph: Paragraph, text: str = "", style: str = "Normal") -> Paragraph:
    new_paragraph = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def paragraph_index(document: Document, target: Paragraph) -> int:
    return next(index for index, paragraph in enumerate(document.paragraphs) if paragraph._p is target._p)


def add_diagram(document: Document, anchor: Paragraph) -> None:
    if any(CAPTION in paragraph.text for paragraph in document.paragraphs):
        return
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(DIAGRAM), width=Inches(6.4))
    caption = document.add_paragraph(CAPTION)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        caption.style = document.styles["Caption"]
    except KeyError:
        pass
    anchor._p.addnext(picture._p)
    picture._p.addnext(caption._p)


def update_verified_test_count(document: Document) -> None:
    replacements = {
        "6,567": "6,332",
        "6,435": "6,332",
        "6,340": "6,332",
        "5,619": "5,362",
        "5,465": "5,362",
        "5,370": "5,362",
        "XAML 490": "XAML 512",
        "src 4,965": "src 4,865",
        "src 4,968": "src 4,865",
        "src 4,873": "src 4,865",
        "tests 1,118": "tests 983",
        "支援コード1,499行": "支援コード1,713行",
        "支援コード1,713行": "支援コード1,719行",
        "合計は7,831行": "合計は8,051行",
        "合計は7,831": "合計は8,051",
        "合計は8,045行": "合計は8,051行",
        "合計は8,045": "合計は8,051",
        "WPF・WinUI・RecognizerのARM64/x64": "WPF・WinUIのARM64/x64",
        "ARM64 MSIX 32,049,628 bytes": "ARM64 MSIX 0.6.0.4 32,048,102 bytes",
        "起動時確認推論9.250 ms、デモ筆跡推論1.542 ms": "デモ筆跡20回の中央値1.524 ms、P95 5.400 ms",
        "WPF 0.6.0 MSIXの起動時確認推論は9.250 ms、デモ筆跡推論は1.542 ms": "WPF 0.6.0.4 ARM64 MSIXのデモ筆跡20回は中央値1.524 ms、P95 5.400 ms",
        "NotPresent→NotReady→Ready": "NotPresentまたはNotReady→Ready",
        "初回の単一計測": "同一端末20回の計測",
    }
    for paragraph in document.paragraphs:
        if "73/73" in paragraph.text or "73件" in paragraph.text:
            paragraph.text = paragraph.text.replace("73/73", "66/66").replace("73件", "66件")
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
        if "物理LOC" in paragraph.text and "6,332行" in paragraph.text and "8,051行" not in paragraph.text:
            paragraph.text += " 支援コード1,719行を含むチーム作成コードの合計は8,051行である。"
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "73/73" in paragraph.text or "73件" in paragraph.text:
                        paragraph.text = paragraph.text.replace("73/73", "66/66").replace("73件", "66件")
                    for old, new in replacements.items():
                        if old in paragraph.text:
                            paragraph.text = paragraph.text.replace(old, new)


def update_plan() -> None:
    document = Document(PLAN)
    update_verified_test_count(document)
    structure = find(document, "8.1 構成", last=True)
    body_index = paragraph_index(document, structure)
    following = document.paragraphs[body_index + 1]
    add_diagram(document, following)

    if not any("8.5 実装規模" in paragraph.text for paragraph in document.paragraphs):
        anchor = find(document, "開発支援AIと実行時AIは役割", last=True)
        heading = insert_after(anchor, "8.5 実装規模とチーム実装", "Heading 2")
        paragraph = insert_after(
            heading,
            "生成物、bin、obj、外部ライブラリを除くアプリ・テスト・学習実験の物理LOCは6,332行"
            "（C# 5,362、XAML 512、Python 458）である。内訳はsrc 4,865、tests 983、experiments 484である。"
            "インストール、実機QA、公開監査、文書生成、デモ録画の支援コードは別に1,719行あり、合計は8,051行である。"
            "空行とコメントを含む。チーム実装は、"
            "ペン入力UI、記録・固定観測、本人意思優先制御、保存制御、軌跡モデルの学習・量子化、Windows ML/QNN選択、"
            "テストである。OSの文字候補・音声、Windows ML、QNN、第三者ライブラリとKanjiVGは外部資源として区別する。",
        )
        heading2 = insert_after(paragraph, "8.6 実測指標と参照資料", "Heading 2")
        insert_after(
            heading2,
            "Core単体テスト66/66、WPF・WinUIのARM64/x64 Releaseビルド警告0・エラー0、"
            "ARM64 MSIX 0.6.0.4は32,048,102 bytes、QDQ INT8モデルは115,925 bytesである。Surface Pro 11では"
            "QNNExecutionProvider / Qualcomm Hexagon NPUを明示選択し、CPUフォールバックを無効にした。"
            "デモ筆跡20回は中央値1.524 ms、P95 5.400 msで、Windows日本語認識と本人意思優先フローも回帰確認した。"
            "最大ワーキングセット1,087.2 MiBは制約として扱う。詳細と適用境界は「WriteMirror_実装・性能指標_ja.md」、外部資源と"
            "ライセンスは「WriteMirror_プラットフォーム・データ・ライセンス_ja.md」に記録する。",
        )

    document.core_properties.author = "WriteMirror Project"
    document.core_properties.last_modified_by = "WriteMirror Project"
    temporary = PLAN.with_name(PLAN.stem + ".tmp.docx")
    document.save(temporary)
    temporary.replace(PLAN)


def update_guide() -> None:
    document = Document(GUIDE)
    update_verified_test_count(document)
    architecture = find(document, "9. 技術構成", last=True)
    index = paragraph_index(document, architecture)
    following = document.paragraphs[index + 1]
    add_diagram(document, following)

    ai_heading = find(document, "AI・NPU版（0.6.0）の技術説明", last=True)
    ai_heading.text = "16. AI・NPU版（0.6.0）の技術説明"

    if not any("17. 実装量・外部資源" in paragraph.text for paragraph in document.paragraphs):
        anchor = document.paragraphs[-1]
        heading = insert_after(anchor, "17. 実装量・外部資源・性能の確認先", "Heading 1")
        insert_after(
            heading,
            "アプリ・テスト・学習実験の物理LOCは6,332行、支援コードを含む合計は8,051行である。"
            "チーム実装と外部プラットフォーム、データセット、ライセンスの区別は"
            "「WriteMirror_プラットフォーム・データ・ライセンス_ja.md」、モデルサイズ、テスト結果、実機NPU値、"
            "オフライン性は「WriteMirror_実装・性能指標_ja.md」を参照する。発表スライドと発表原稿は他の"
            "グループメンバーが担当し、本説明書と実機結果との整合を確認する。",
        )

    document.core_properties.author = "WriteMirror Project"
    document.core_properties.last_modified_by = "WriteMirror Project"
    temporary = GUIDE.with_name(GUIDE.stem + ".tmp.docx")
    document.save(temporary)
    temporary.replace(GUIDE)


def main() -> None:
    update_plan()
    update_guide()
    print(PLAN)
    print(GUIDE)


if __name__ == "__main__":
    main()
