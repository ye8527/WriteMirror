"""Synchronize the Japanese DOCX files with the verified Windows ML implementation."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


DOCS = Path(__file__).resolve().parents[1] / "docs"


def insert_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_paragraph = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def set_cell(table, row: int, column: int, text: str) -> None:
    table.rows[row].cells[column].text = text


def save(document: Document, path: Path) -> None:
    document.core_properties.author = "WriteMirror Project"
    document.core_properties.last_modified_by = "WriteMirror Project"
    temporary = path.with_name(path.stem + ".tmp.docx")
    document.save(temporary)
    temporary.replace(path)


def update_plan(path: Path) -> None:
    document = Document(path)
    paragraphs = document.paragraphs
    if any("Windows ML・Qualcomm NPUへの実行時AI統合" in paragraph.text for paragraph in paragraphs):
        for paragraph in paragraphs:
            paragraph.text = paragraph.text.replace(
                "起動時の単一確認推論は213.251 ms",
                "起動時確認推論は初回213.251 ms、再起動後7.979 ms",
            )
            paragraph.text = paragraph.text.replace(
                "起動時確認推論は初回213.251 ms、再起動後7.979 ms",
                "WPF 0.6.0.4 ARM64 MSIXのデモ筆跡20回は中央値1.524 ms、P95 5.400 ms",
            )
        set_cell(
            document.tables[8],
            5,
            2,
            "WPF 0.6.0 ARM64 MSIX。Surface Pro 11のHexagon NPUで確認推論済み",
        )
        if not any("観測を見てみる" in paragraph.text and "AI推論" in paragraph.text for paragraph in paragraphs):
            anchor = next(
                paragraph
                for paragraph in paragraphs
                if "モデルが返す再構成差" in paragraph.text
            )
            insert_after(
                anchor,
                "本人意思優先の回帰確認では、「特になし」の確定直後は入力別AI推論0件、本人が"
                "「観測を見てみる（任意）」を押した後に1件となった。AI値は本人回答を先回りして表示しない。",
            )
        save(document, path)
        return
    assert "独自AIモデル" in paragraphs[88].text
    assert "生成AIをコード案" in paragraphs[110].text
    assert "現在の実験モデル" in paragraphs[115].text
    assert "実行時AIを追加しない判断" in paragraphs[122].text

    paragraphs[88].text = (
        "Phi SilicaまたはAion Instructが実接続済みであるという主張、およびAIによる診断・能力評価"
    )
    paragraphs[110].text = (
        "本開発では、生成AIをコード案、テスト案、文書草案、レビュー観点の作成に利用した。"
        "生成物はそのまま正しいとはみなさず、ビルド、単体テスト、実機確認、文献照合によって検証した。"
        "これとは別に、0.6.0開発版の完成アプリはWindows ML上で独自の量子化ONNX軌跡モデルを実行する。"
        "開発支援AIと実行時AIは役割と検証証拠を分けて説明する。"
    )
    paragraphs[115].text = (
        "比較実験として、KanjiVGの整形済み軌跡を扱う双方向GRUを保持する。加えて、QNN HTPの静的形状・"
        "量子化要件に合わせ、128点×3特徴量を入力する全結合オートエンコーダーを実装した。構成は"
        "384→128→32→128→384であり、正規化座標と筆画末尾から軌跡を再構成する。児童の時刻、ペン先圧、"
        "傾き、困難、能力、診断、感情または改善ラベルは学習していない。"
    )
    paragraphs[122].text = "9.4 Windows ML・Qualcomm NPUへの実行時AI統合"
    paragraphs[123].text = (
        "0.6.0開発版では、Windows MLのExecutionProviderCatalogを用いて認定済みQNN実行プロバイダーの"
        "ReadyStateを確認し、NotPresentまたはNotReadyの場合はEnsureReadyAsyncを呼び出す。これにより、"
        "開発者がQNNバイナリを個別管理せず、OS管理で取得・更新できる。登録後はONNX Runtimeから"
        "QNNExecutionProviderのNPUデバイスを明示選択する。"
    )
    first = insert_after(
        paragraphs[123],
        "Surface Pro 11（Snapdragon X Elite）の実機では、NotPresent→NotReady→Readyの遷移、認定QNN EPの"
        "取得・登録、CPUフォールバックを無効にした量子化ONNXセッション、確認推論まで成功した。"
        "起動時の単一確認推論は213.251 msであり、性能比較値ではなく技術経路の実行証拠として記録する。",
    )
    insert_after(
        first,
        "モデルが返す再構成差は研究上の中立的観測値であり、文字の正しさ、良否、読みやすさ、困難、能力、"
        "診断または教育効果を意味しない。本人の回答を上書きせず、NPUが利用できない場合はCPUモードと明示する。"
        "Phi SilicaはLimited Access Featureと移行時期の制約があるため実装済みとはしない。",
    )

    table = document.tables[8]
    set_cell(table, 5, 0, "AI / NPU")
    set_cell(table, 5, 1, "Windows ML readiness、認定QNN EP、量子化ONNX軌跡モデル")
    set_cell(table, 5, 2, "0.6.0開発版。Surface Pro 11のHexagon NPUで確認推論済み")

    table = document.tables[9]
    set_cell(table, 4, 0, "軌跡AI・NPU")
    set_cell(table, 4, 1, "KanjiVG 11,662件で学習し、QDQ INT8 ONNXをQNN NPUで確認推論")
    set_cell(table, 4, 2, "診断・能力評価・文字正誤・児童への有効性")

    table = document.tables[16]
    set_cell(table, 6, 0, "AI/NPUの誤表示")
    set_cell(table, 6, 1, "実行プロバイダーやフォールバックの説明が事実と不一致になる")
    set_cell(table, 6, 2, "EP名、デバイス種別、推論時間をログ化し、CPU時はNPUと表示しない")

    table = document.tables[17]
    set_cell(table, 4, 0, "量子化軌跡モデルをアプリへ統合し、QNN NPUで確認推論した")
    set_cell(table, 4, 1, "Windows MLによるOS管理のreadinessとオンデバイス推論を提示できる")
    set_cell(table, 4, 2, "児童の困難・能力・正しさ・改善または教育効果を推論できる")

    document.add_paragraph(
        "[18] Microsoft. Install Windows ML execution providers. "
        "https://learn.microsoft.com/en-us/windows/ai/new-windows-ml/initialize-execution-providers"
    )
    document.add_paragraph(
        "[19] Microsoft. Select execution providers using the ONNX Runtime included in Windows ML. "
        "https://learn.microsoft.com/en-us/windows/ai/new-windows-ml/select-execution-providers"
    )
    document.add_paragraph(
        "[20] ONNX Runtime. QNN Execution Provider. "
        "https://onnxruntime.ai/docs/execution-providers/QNN-ExecutionProvider.html"
    )
    save(document, path)


def update_manual(path: Path) -> None:
    document = Document(path)
    if any("AI・NPU開発版" in paragraph.text for paragraph in document.paragraphs):
        for paragraph in document.paragraphs:
            paragraph.text = paragraph.text.replace("AI・NPU開発版（0.6.0）", "AI・NPU版（0.6.0）")
            paragraph.text = paragraph.text.replace("0.6.0開発版は", "0.6.0は")
            paragraph.text = paragraph.text.replace(
                "起動時確認推論213.251 ms",
                "起動時確認推論は初回213.251 ms、再起動後7.979 ms",
            )
            paragraph.text = paragraph.text.replace(
                "起動時確認推論は初回213.251 ms、再起動後7.979 ms",
                "WPF 0.6.0.4 ARM64 MSIXのデモ筆跡20回は中央値1.524 ms、P95 5.400 ms",
            )
        if not any("観測を見てみる" in paragraph.text and "AI推論" in paragraph.text for paragraph in document.paragraphs):
            document.add_paragraph(
                "本人意思優先の確認では、「特になし」の確定直後はAI推論を表示せず、本人が"
                "「観測を見てみる（任意）」を押した後だけ入力別のモデル結果を表示します。"
            )
        save(document, path)
        return
    table = document.tables[0]
    assert "独自AI" in table.rows[5].cells[1].text
    set_cell(table, 5, 1, "診断、治療、能力評価、教育効果判定、Phi Silica、Aion Instruct。")

    table = document.tables[7]
    set_cell(table, 4, 0, "公開KanjiVGで学習した量子化軌跡モデルをアプリへ統合した。")
    set_cell(table, 4, 1, "Windows MLの認定QNN EPを介し、SurfaceのHexagon NPUで実行する。")
    set_cell(table, 4, 2, "児童の困難、能力、正しさ、改善、診断または教育効果を推論する。")

    document.add_heading("AI・NPU開発版（0.6.0）の技術説明", level=1)
    document.add_paragraph(
        "0.6.0開発版は、Windows MLのExecutionProviderCatalogでQNN実行プロバイダーの準備状態を確認し、"
        "必要な場合にEnsureReadyAsyncを呼び出します。Surface Pro 11では、認定QNN EPの取得、登録、"
        "Hexagon NPUデバイスの明示選択、量子化ONNXモデルの推論まで確認しました。"
    )
    for text in (
        "モデル：KanjiVG 11,662件で自己教師あり学習した固定長軌跡オートエンコーダー。",
        "入力：128点×正規化X・正規化Y・筆画終端。実測の困難・診断・能力ラベルは使用しません。",
        "実行：ARM64 SurfaceではQNN NPUを優先し、CPUフォールバックを無効にした確認セッションを使用します。",
        "表示：モデル名、実行プロバイダー、推論時間、軌跡再構成差を表示します。",
        "制約：再構成差は採点、診断、能力評価、良否判定または本人の回答の上書きに使用しません。",
        "代替：NPU非搭載または準備失敗時はCPUモードと明示し、基本の書字記録機能を継続します。",
    ):
        document.add_paragraph(text, style="List Bullet")

    document.add_paragraph(
        "実機確認値：QNNExecutionProvider / Qualcomm Hexagon NPU、起動時確認推論213.251 ms。"
        "これは初回の単一計測であり、性能比較や製品保証には使用しません。"
    )
    save(document, path)


def main() -> None:
    plan = next(DOCS.glob("*Vibe_Coding*.docx"))
    manual = next(DOCS.glob("*ja.docx"))
    update_plan(plan)
    update_manual(manual)
    print(plan)
    print(manual)


if __name__ == "__main__":
    main()
