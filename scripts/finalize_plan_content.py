"""Finalize the research plan content for the verified WriteMirror 0.6.0 scope."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
PLAN = next((ROOT / "docs").glob("*Vibe_Coding*.docx"))


def find(document: Document, fragment: str) -> Paragraph:
    return next(paragraph for paragraph in document.paragraphs if fragment in paragraph.text)


def insert_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_paragraph = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def set_cell(table, row: int, column: int, text: str) -> None:
    table.rows[row].cells[column].text = text


def main() -> None:
    document = Document(PLAN)

    problem = find(document, "以上から、本研究の中心課題")
    if not any("単項目の規則" in paragraph.text for paragraph in document.paragraphs):
        insert_after(
            problem,
            "AIが必要となる理由は、児童を自動判定するためではない。総時間や筆画数などの単項目の規則は"
            "個別の事実を説明しやすい一方、128点から成る複数筆画の位置関係、順序、全体形状を一つの"
            "表現として扱えない。本研究の軌跡オートエンコーダーは、公開軌跡から全体パターンを圧縮・再構成し、"
            "規則ベースの観測を補う任意の参考表示を作る。筆跡をクラウドへ送らず端末内NPUで処理することで、"
            "個人性を含み得る筆跡の外部送信を避ける。",
        )

    rq5 = find(document, "RQ5：")
    if not any("RQ6：" in paragraph.text for paragraph in document.paragraphs):
        insert_after(
            rq5,
            "RQ6：単項目の観測値だけでは扱えない複数筆画の全体軌跡を、オンデバイスAIで圧縮・再構成し、"
            "本人の回答を上書きしない任意表示として提示できるか。",
            "List Bullet",
        )

    h4 = find(document, "H4：")
    if not any("H5：" in paragraph.text for paragraph in document.paragraphs):
        insert_after(
            h4,
            "H5：本人が観測を希望した後に限って原軌跡とAI再構成を並べると、単独の数値だけよりも"
            "書字全体を振り返りやすくなる。ただし正解見本または良否判定とは扱わない。",
            "List Bullet",
        )

    scenario_anchor = find(document, "研究・技術教育：")
    if not any("想定利用場面の例" in paragraph.text for paragraph in document.paragraphs):
        scenario_heading = insert_after(scenario_anchor, "6.3 想定利用場面の例", "Heading 2")
        scenario = insert_after(
            scenario_heading,
            "小学2年生の利用者が授業中に「木」を練習する場面を想定する。本人が「特になし」を選んだ場合、"
            "システムは問題候補やAI値を自動表示せず、その回答を受け取る。本人が自分から「観測を見てみる」"
            "を選んだ場合だけ、書いた順序、画間の観測、原軌跡とAI再構成を表示する。本人は正誤判定を受けず、"
            "必要ならもう一度書く。教師は診断結果ではなく、本人との会話を始める材料として画面を見る。",
        )
        next_heading = find(document, "6.3 本研究で扱わない場面")
        next_heading.text = "6.4 本研究で扱わない場面"

    heading = find(document, "9.3 本研究でモデルを追加訓練しない理由")
    heading.text = "9.3 モデル学習の範囲と適用境界"
    boundary = find(document, "9.4 Windows ML・Qualcomm NPU")
    between: list[Paragraph] = []
    started = False
    for paragraph in document.paragraphs:
        if paragraph._p is heading._p:
            started = True
            continue
        if paragraph._p is boundary._p:
            break
        if started and paragraph.text.strip():
            between.append(paragraph)
    replacements = [
        "本課題ではKanjiVG 11,662件を用いた自己教師あり学習を実施し、固定長軌跡オートエンコーダーを作成した。",
        "学習対象は正規化座標と筆画終端であり、困難、障害、能力、正誤、感情または改善ラベルを含まない。",
        "KanjiVGは整形済みの漢字軌跡であり、低学年児童の実測筆跡、平仮名、片仮名の妥当性を代表しない。",
        "児童データを用いた追加学習、個人適応、診断モデル化、大規模モデル開発は本課題では行わない。",
        "モデル出力は原軌跡の圧縮・再構成を観察する研究値に限定し、Windows文字認識や本人の回答を置き換えない。",
    ]
    for paragraph, text in zip(between, replacements, strict=False):
        paragraph.text = text
        paragraph.style = document.styles["List Bullet"]
    for paragraph in between[len(replacements):]:
        paragraph._element.getparent().remove(paragraph._element)

    for paragraph in document.paragraphs:
        paragraph.text = paragraph.text.replace("0.6.0開発版", "0.6.0")
        paragraph.text = paragraph.text.replace("WriteMirror 0.5.2 ARM64実行版", "WriteMirror 0.6.0 ARM64 MSIX実行版")
        paragraph.text = paragraph.text.replace("90秒実機デモ", "0.6.0実機デモ動画")

    role_table = document.tables[14]
    if not any("発表スライド" in row.cells[0].text for row in role_table.rows):
        row = role_table.add_row()
        row.cells[0].text = "発表スライド・発表講稿"
        row.cells[1].text = "他のグループメンバーが担当"
        row.cells[2].text = "本計画書・実機結果との整合を確認"

    stage_table = document.tables[15]
    for row in stage_table.rows:
        label = row.cells[0].text
        if label.startswith("3."):
            row.cells[1].text = "WPF ARM64、Surface Pen、音声、文字候補、保存方針、軌跡AI"
            row.cells[2].text = "完了"
        elif label.startswith("5."):
            row.cells[1].text = "Coreテスト、ARM64 MSIX、Windows ML readiness、QNN NPU推論"
            row.cells[2].text = "完了"
        elif label.startswith("6."):
            row.cells[1].text = "研究計画書、成果説明、0.6.0デモ動画、GitHub公開"
            row.cells[2].text = "デモ動画・GitHub更新中"

    document.core_properties.author = "WriteMirror Project"
    document.core_properties.last_modified_by = "WriteMirror Project"
    temporary = PLAN.with_name(PLAN.stem + ".tmp.docx")
    document.save(temporary)
    temporary.replace(PLAN)
    print(PLAN)


if __name__ == "__main__":
    main()
