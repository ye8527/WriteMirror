from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "WriteMirror_利用・技術説明書_ja.docx"
VALIDATION_PATH = DOCS / "WriteMirror_検証結果_ja.md"
REFERENCES_PATH = DOCS / "WriteMirror_先行研究・参考文献_ja.md"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = label
        set_cell_shading(cell, "DCE6F1")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


DOCS.mkdir(parents=True, exist_ok=True)
doc = Document()
doc.core_properties.title = "WriteMirror 利用・技術説明書"
doc.core_properties.subject = "WriteMirror 0.5.2の利用方法、技術構成、検証範囲"
doc.core_properties.author = "WriteMirror Project"
doc.core_properties.last_modified_by = "WriteMirror Project"
doc.core_properties.comments = "日本語版・公開用メタデータ"
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = doc.styles
styles["Normal"].font.name = "Yu Gothic"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
styles["Normal"].font.size = Pt(10.5)
for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
    styles[style_name].font.name = "Yu Gothic"
    styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("WriteMirror\n利用・技術説明書")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(31, 78, 121)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("小グループ研究課題用プロトタイプ 0.5.2\n2026年7月18日").font.size = Pt(12)

doc.add_paragraph()
notice = doc.add_paragraph()
notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
notice_run = notice.add_run("教育効果・診断精度・児童だけでの安全な利用を実証した製品ではありません。")
notice_run.bold = True
notice_run.font.color.rgb = RGBColor(192, 0, 0)

doc.add_heading("1. 文書の目的", level=1)
doc.add_paragraph(
    "本書は、WriteMirror 0.5.2の利用方法、技術構成、対応環境、データ取扱い、検証範囲と制約を、"
    "発表者・評価者・利用者が同じ意味で理解できるようにまとめたものです。正式な研究背景、先行研究、"
    "研究課題および評価計画は、別冊の「WriteMirror_研究計画書_ja.docx」を参照してください。"
)

doc.add_heading("2. アプリの位置づけ", level=1)
add_table(doc, ["項目", "内容"], [
    ["目的", "完成文字を採点せず、書いている途中の筆跡と本人の振り返りを並べて確認する。"],
    ["想定する対象像", "書字につまずきのある可能性を含む児童・小学校低学年。ただし本課題では児童を研究参加者にしない。"],
    ["利用場面", "授業内の成人による技術デモ、研究計画と実装可能性の説明。"],
    ["実装上の独自性", "本人の回答を先に受け取り、希望した場合だけ端末の観測候補を提示し、本人が再試行を選べる流れ。"],
    ["実装していないもの", "診断、治療、能力評価、教育効果判定、独自AI、Phi Silica、Aion Instruct、NPU推論。"],
])

doc.add_heading("3. 対応環境", level=1)
add_table(doc, ["配布版", "想定端末", "必要条件"], [
    ["ARM64", "Snapdragon X 系などのWindows 11 ARM64 Surface", "Windows 11、Surface Penまたは互換ペン。日本語文字候補にはWindowsの日本語手書き認識機能が必要。"],
    ["x64", "IntelまたはAMD搭載のWindows 11 Surface／PC", "Windows 11、Windows Ink対応ペン。日本語文字候補にはWindowsの日本語手書き認識機能が必要。"],
])
doc.add_paragraph(
    "本アプリは.NETランタイムを同梱した自己完結型です。通常は追加の.NETインストールを必要としません。"
    "端末固有のペン機能、Windowsの言語機能、組織のアプリ制御方針によって利用できる機能は異なります。"
)

doc.add_heading("4. 起動方法", level=1)
add_numbered(doc, [
    "配布ZIPを端末内の通常フォルダーへすべて展開する。ZIP内から直接実行しない。",
    "最上位にある「WriteMirrorを起動.cmd」をダブルクリックする。",
    "Windowsが確認画面を表示した場合は、発表担当者または端末管理者が配布元とハッシュ値を確認し、組織の方針に従う。",
    "アプリの最初の画面で説明を読み、「つかう」または「いまはやめる」を本人が選ぶ。",
])
doc.add_paragraph(
    "個別起動では、ARM64端末は「app\\ARM64\\WriteMirror.exe」、x64端末は"
    "「app\\x64\\WriteMirror.exe」を使用します。各アプリ配下のRecognizerフォルダーは移動しないでください。"
)

doc.add_heading("5. 基本操作", level=1)
add_numbered(doc, [
    "通常は「ひとりで練習（保存しない）」を選ぶ。共同確認で保存する場合は、目的と削除期限を先に説明する。",
    "平仮名・片仮名・漢字など、書きたい文字を自由に入力する。Windowsの日本語文字候補が表示された場合は、本人が候補を選べる。",
    "Surface Penで枠内に書く。必要に応じて消去または書き直しを行う。",
    "分析を見る前に、「ためらった」「書きにくかった」「気になった」「特になし」「うまくいった」「答えない」から選ぶ。",
    "位置を伴う回答を選んだ場合だけ、気になった範囲を囲む。",
    "「特になし」「うまくいった」「答えない」の場合、観測候補は自動表示されない。本人が「観測を見てみる」または「これで終わる」を選ぶ。",
    "本人が希望する場合だけ二回目を書く。数値の増減を良否や能力の変化と解釈しない。",
])

doc.add_heading("6. 画面表示の読み方", level=1)
add_table(doc, ["表示", "意味", "意味しないこと"], [
    ["筆跡再生", "端末が取得した接触中の点列を順に再生する。", "正しい筆順、学習到達度、診断結果。"],
    ["画間空白時間候補", "前の画が終わってから次の画が始まるまでに観測された時間候補。", "ためらい、困難、誤り、原因。"],
    ["筆圧の統計", "端末が取得できた圧力値の範囲や代表値。", "握力、情緒、障害、医学的状態。"],
    ["二回の差", "暫定条件を満たす差が観測されたこと。", "改善、悪化、練習効果。"],
    ["文字候補", "Windowsの日本語手書き認識が返した候補。", "独自AIの推論、100%正しい判定。"],
])

doc.add_heading("7. 本人の意思を優先する設計", level=1)
add_bullets(doc, [
    "本人の回答を分析表示より先に受け取る。",
    "「特になし」「うまくいった」「答えない」を客観候補で上書きしない。",
    "観測候補を見るか、そのまま終えるかを同じ強さのボタンで選べる。",
    "二回目を強制しない。いつでも終了と現在データの削除を選べる。",
    "主画面で点数、順位、赤点、能力ラベルを表示しない。",
])

doc.add_heading("8. 保存と個人情報", level=1)
add_table(doc, ["モード", "保存"], [
    ["ひとりで練習", "保存しない。処理中のメモリーだけで扱う。"],
    ["いっしょに確認", "保存は初期値OFF。本セッションで明示的に選んだ場合だけ端末内へ保存する。"],
])
doc.add_paragraph("保存先：%LOCALAPPDATA%\\WriteMirror\\Sessions")
add_bullets(doc, [
    "保存候補：課題ID、利き手、時刻、接触中の筆画点、取得可能な筆圧・傾き、本人が囲んだ範囲。",
    "保存しないもの：氏名、メールアドレス、学校名、診断情報、生成AIプロンプト。",
    "共同確認で保存する場合は、利用目的、閲覧者、保持期限、削除担当を利用前に決める。",
    "同意を解除した場合は、そのセッションの端末内データを削除する。",
])

doc.add_heading("9. 技術構成", level=1)
add_table(doc, ["構成要素", "役割"], [
    ["WriteMirror.Wpf", "WPF／Windows Inkによる現行UI、ペン入力、再生、音声案内。"],
    ["WriteMirror.Core", "接触点、観測値、比較条件、本人回答ポリシー、固定日本語フィードバック。"],
    ["WriteMirror.Infrastructure", "原子的JSON保存、読込、個別削除、全削除。配布本体では必要部分をWPFへ統合。"],
    ["WriteMirror.Recognizer", "Windowsの日本語手書き認識を別プロセスで呼び出す補助機能。ARM64／x64を分離。"],
    ["WriteMirror.App", "旧WinUI 3試作。実ペン接触時の不安定性があるため最終実演には使用しない。"],
])

doc.add_heading("10. 二回の審査で修正した重要事項", level=1)
add_bullets(doc, [
    "補間時刻を実測局所速度とみなす「低速度区間候補」を画面、照合、音声から削除した。",
    "画間空白時間は未観測の空中経路を線で結ばず、前画終点と次画始点だけを示す。",
    "本人が「特になし」「うまくいった」「答えない」を選んだ場合、候補表示、位置指定、再試行を自動要求しない。",
    "回答種別をフィードバック生成へ渡し、本人の回答を固定テンプレートが上書きしない。",
    "独立練習の保存禁止を画面状態だけでなく核心ポリシーで保証し、原子的JSON保存と削除処理を整備した。",
])

doc.add_heading("11. 検証結果", level=1)
add_table(doc, ["検証", "結果"], [
    ["Core単体テスト", "73件中73件成功。"],
    ["WPF Releaseビルド", "ARM64／x64とも警告0、エラー0。"],
    ["Recognizer Releaseビルド", "ARM64／x64とも警告0、エラー0。"],
    ["WinUI試作ビルド", "ARM64／x64とも警告0、エラー0。ただし最終実演対象外。"],
    ["ARM64実機", "最終候補の起動・応答、日本語認識ヘルパー起動、修正後の主要画面フローを確認。"],
    ["x64実機", "ビルドとx64認識ヘルパーのエミュレーション実行を確認。Intel／AMD Surfaceでの最終操作確認は未実施。"],
])

doc.add_heading("12. トラブル対応", level=1)
add_table(doc, ["症状", "確認事項"], [
    ["起動しない", "ZIPをすべて展開したか、端末の構成に合う版か、組織のアプリ制御が止めていないかを確認する。"],
    ["ペンで書けない", "Surface Penの接続・電池、Windows Ink、対象枠内への接触、別アプリでのペン動作を確認する。"],
    ["日本語候補が出ない", "Windowsの日本語言語機能と手書き認識機能を確認する。対象文字は手入力で継続できる。"],
    ["音声が聞こえない", "端末音量、出力先、ミュート状態を確認する。音声なしでも画面表示で操作できる。"],
    ["保存したデータを消したい", "アプリの削除操作を使うか、共同確認の管理担当者が保存先を確認して削除する。"],
])

doc.add_heading("13. 発表で述べてよい範囲", level=1)
add_table(doc, ["述べてよいこと", "条件付きで述べること", "述べないこと"], [
    ["Surface Penで書字過程を記録・再生できる。", "Windowsが日本語文字候補を返す。", "すべての文字を100%認識できる。"],
    ["本人の回答と観測位置を並べられる。", "二回の値に差が観測された。", "書字能力が改善した。"],
    ["独立練習は保存しない設計である。", "児童の独立利用を想定した候補UIである。", "児童だけで安全に使えることを実証した。"],
    ["公開資料を用いた独立実験がある。", "将来は別方式のモデルを検討できる。", "独自AI、Phi Silica、Aion、NPUが動作している。"],
])

doc.add_heading("14. 残る制約と今後の研究課題", level=1)
add_bullets(doc, [
    "児童を対象とする有効性、理解可能性、無介助完遂、安全性は未検証である。",
    "端末・ペン・サンプリング方式の違いによる観測値の再現性は未検証である。",
    "教育的効果、学校運用、家庭運用、支援者の負担軽減は研究仮説である。",
    "公開データだけで児童の困難を推論しない。将来の学習には目的適合性、利用許諾、児童データ倫理、外部検証が必要である。",
    "別端末での無条件の動作保証は行わない。導入前に対象端末と組織ポリシーで受入確認を行う。",
])

doc.add_heading("15. 配布物", level=1)
add_bullets(doc, [
    "app：ARM64版、x64版、構成自動判定ランチャー。",
    "documents：研究計画書、本説明書、検証結果、先行研究一覧、README、ファイル一覧。",
    "demo：日本語ナレーション付きMP4。",
    "source：実装、単体テスト、文書生成・デモ生成に必要なスクリプト。",
    "整合性確認：SHA-256一覧。",
])

doc.save(DOCX_PATH)

VALIDATION_PATH.write_text("""# WriteMirror 0.5.2 検証結果

検証日：2026年7月18日

## 結果

| 対象 | 結果 |
|---|---|
| Core単体テスト | 73/73 成功 |
| WPF ARM64 / x64 Release | 警告0、エラー0 |
| Recognizer ARM64 / x64 Release | 警告0、エラー0 |
| WinUI ARM64 / x64 Release | 警告0、エラー0（旧試作、最終実演対象外） |
| ARM64最終候補 | Surface実機で起動・応答を確認 |
| ARM64日本語認識ヘルパー | 実行とJSON応答を確認 |
| x64日本語認識ヘルパー | Windows on ARMのx64エミュレーションで実行とJSON応答を確認 |

## P0修正の確認

- 補間時刻に基づく低速度・最遅区間候補は、画面、照合、音声から削除した。
- 「特になし」「うまくいった」「答えない」では候補を自動表示しない。
- 上記3回答では位置指定と再試行を自動要求しない。
- 本人が「観測を見てみる」を選んだ場合だけ、画間空白時間候補を表示する。
- 独立練習では保存処理を核心ポリシーが拒否する。

## 実機確認の範囲

ARM64版は現在のSurface実機で主要フローを確認した。x64版は同一ソースから自己完結形式で生成したが、現在のARM64端末では組織のアプリ制御がx64 GUI本体を停止したため、IntelまたはAMD搭載Surfaceでの最終操作確認は未実施である。別端末での無条件の動作保証を意味しない。

## 未検証事項

- 児童だけでの利用、安全性、理解可能性、教育的効果
- 学校または家庭での継続運用
- 診断、治療、能力評価
- 独自AI、Phi Silica、Aion Instruct、NPU推論
- すべての平仮名・片仮名・漢字に対する認識精度
""", encoding="utf-8")

REFERENCES_PATH.write_text("""# WriteMirror 先行研究・参考文献

正式な研究背景、各文献の位置づけ、研究課題への反映は「WriteMirror_研究計画書_ja.docx」に記載する。ここでは配布時の確認用として主要資料を一覧化する。

1. Rosenblum, S., Weiss, P. L., & Parush, S. (2003). Product and process evaluation of handwriting difficulties. *Educational Psychology Review*. https://doi.org/10.1023/A:1021371425220
2. Corbillé, S., Fromont, É., Anquetil, É., & Nerdeux, P. (2020). Integrating Writing Dynamics in CNN for Online Children Handwriting Recognition. *ICFHR 2020*. https://doi.org/10.1109/ICFHR2020.2020.00057
3. Inoue, T., Chen, Y., & Ohyanagi, T. (2024). Assessing handwriting skills in a web browser: Development and validation of an automated online test in Japanese Kanji. *Behavior Research Methods*. https://doi.org/10.3758/s13428-024-02562-6
4. Alamargot, D., Morin, M.-F., & Simard-Dupuis, É. (2015). Does handwriting on a tablet screen affect students’ graphomotor execution? *Human Movement Science, 44*, 32–41. https://doi.org/10.1016/j.humov.2015.08.011
5. Hochhauser, M., Wagner, M., & Shvalb, N. (2023). Assessment of children's writing features: A pilot method study of pen-grip kinetics and writing surface pressure. *Assistive Technology, 35*(1), 107–115. https://doi.org/10.1080/10400435.2021.1956640
6. Nakamoto, R., Flanagan, B., Nakamura, K., & Ogata, H. (2026). Explain-from-Stroke: Capturing Invisible Learning Processes Through Handwriting Dynamics Analysis. *AAAI 2026*. https://doi.org/10.1609/aaai.v40i48.42118
7. 産業技術総合研究所. ETL文字データベース. https://etlcdb.db.aist.go.jp/the-etl-character-database/
8. 東京農工大学 中川研究室. HANDS-nakayosi_t-98-09. https://web.tuat.ac.jp/~nakagawa/database/en/about_nakayosi.html
9. KanjiVG Project. Kanji Vector Graphics. https://kanjivg.tagaini.net/
10. Microsoft. Pen interactions and Windows Ink in Windows apps. https://learn.microsoft.com/en-us/windows/uwp/ui-input/pen-and-stylus-interactions
11. Microsoft. Phi SilicaからAion Instructへの移行案内. https://learn.microsoft.com/en-us/windows/ai/apis/phi-silica

公開データセットの存在は、児童の困難を推論できることや教育効果を証明するものではない。利用時にはライセンス、目的適合性、対象集団、端末差、評価設計を個別に確認する。
""", encoding="utf-8")

print(DOCX_PATH)
print(VALIDATION_PATH)
print(REFERENCES_PATH)
