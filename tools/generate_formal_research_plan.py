from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "WriteMirror_アプリ開発計画書_Vibe_Coding版.docx"

BLUE = "1F4E79"
MID_BLUE = "2F75B5"
LIGHT_BLUE = "DDEBF7"
PALE_BLUE = "EAF2F8"
PALE_YELLOW = "FFF2CC"
PALE_GREEN = "E2F0D9"
LIGHT_GRAY = "F2F2F2"
TEXT = (45, 45, 45)


def set_run_font(run, size=10.5, bold=False, color=TEXT, name="Yu Gothic UI"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text, *, bold=False, color=TEXT, size=9.2, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade(cell, fill)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, color=(255, 255, 255), fill=BLUE)
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row_no, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, fill=PALE_BLUE if row_no % 2 else None)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items, *, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        set_run_font(p.add_run(item))


def add_note(doc, title, text, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}：")
    set_run_font(r, size=9.5, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9.5)
    doc.add_paragraph()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(paragraph.add_run("Page "), size=8, color=(100, 100, 100))
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_flow(doc):
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        "文字を選ぶ", "→", "ペンで書く", "→", "軌跡を再生する",
        "→", "観察と音声を確認", "→", "もう一度試す",
    ]
    for i, label in enumerate(labels):
        fill = LIGHT_BLUE if i % 2 == 0 else LIGHT_GRAY
        set_cell(table.cell(0, i), label, bold=i % 2 == 0, size=8.7, fill=fill)
    doc.add_paragraph()


def add_reference(doc, number, citation, url):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.left_indent = Cm(0.7)
    set_run_font(p.add_run(f"[{number}] {citation} "), size=9.2)
    r = p.add_run(url)
    set_run_font(r, size=9.2, color=(5, 99, 193))


doc = Document()
doc.core_properties.title = "WriteMirror アプリ開発・研究計画書"
doc.core_properties.subject = "Surface Penを用いた書字プロセス振り返り支援の小グループ研究"
doc.core_properties.author = "WriteMirror プロジェクトチーム"
doc.core_properties.comments = "先行研究統合・授業内小グループ研究版"

section = doc.sections[0]
section.top_margin = Cm(1.9)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

normal = doc.styles["Normal"]
normal.font.name = "Yu Gothic UI"
normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Yu Gothic UI")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.18

for name, size, color in [
    ("Title", 28, (31, 78, 121)),
    ("Heading 1", 17, (31, 78, 121)),
    ("Heading 2", 13.5, (47, 117, 181)),
    ("Heading 3", 11.5, (68, 68, 68)),
]:
    style = doc.styles[name]
    style.font.name = "Yu Gothic UI"
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Yu Gothic UI")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(*color)

if "Caption JP" not in doc.styles:
    caption = doc.styles.add_style("Caption JP", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Yu Gothic UI"
    caption._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Yu Gothic UI")
    caption.font.size = Pt(9)
    caption.font.italic = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(header.add_run("WriteMirror｜アプリ開発・研究計画書"), size=8, color=(100, 100, 100))
add_page_number(section.footer.paragraphs[0])

# Cover
doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(title.add_run("WriteMirror"), size=30, bold=True, color=(31, 78, 121))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    subtitle.add_run("Surface Penを用いた書字プロセス振り返り支援アプリの研究開発"),
    size=16, bold=True, color=(55, 55, 55),
)

edition = doc.add_paragraph()
edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(edition.add_run("アプリ開発・研究計画書（先行研究統合版）"), size=14, color=(47, 117, 181))

doc.add_paragraph()
meta = doc.add_table(rows=7, cols=2)
meta.style = "Table Grid"
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
metadata = [
    ("研究形態", "授業内の小グループ研究・プロトタイプ開発"),
    ("研究対象", "書字プロセス振り返り支援の技術的実現可能性"),
    ("想定利用者", "書字に困り感を持つ可能性のある児童、小学校低学年、一般学習者"),
    ("実際の評価範囲", "小グループメンバーおよび成人による機能確認・実機デモ"),
    ("対象端末", "Microsoft Surface Pro 11 / ARM64 / Surface Slim Pen 2"),
    ("試作版", "WriteMirror 0.5.2"),
    ("作成日", "2026年7月18日"),
]
for row, (key, value) in enumerate(metadata):
    set_cell(meta.rows[row].cells[0], key, bold=True, fill=LIGHT_BLUE)
    set_cell(meta.rows[row].cells[1], value)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(
    tagline.add_run("診断・採点ではなく、本人の気づきと再試行を支える研究試作"),
    size=11, bold=True, color=(127, 96, 0),
)
doc.add_page_break()

add_heading(doc, "文書の位置づけ", 1)
add_para(
    doc,
    "本書は、WriteMirrorを授業内の小グループ研究として整理した正式な研究計画書兼アプリ技術計画書である。研究背景、先行研究、課題、目的、対象者、利用場面、提案手法、実装、評価、倫理、期待される効果および限界を一貫した形で記述する。",
)
add_para(
    doc,
    "本研究は、医療機器開発、障害診断、児童を対象とした介入研究、学校・家庭への正式導入を目的としない。実児童から新たなデータを収集する計画も置かない。児童は設計上の想定利用者であり、本授業課題で実施する評価は、小グループメンバーおよび成人協力者による技術的・操作的確認に限定する。",
)
add_note(
    doc,
    "重要な解釈",
    "本書に記載する教育的効果は先行研究と設計仮説に基づく「期待される効果」であり、WriteMirrorによって実証済みの効果ではない。",
)

add_heading(doc, "要旨", 1)
add_para(
    doc,
    "書字の評価は、完成した文字の正誤や読みやすさに集中しやすい。しかし、書いている途中の停滞、書き直し、筆画間の空白時間、本人が気になった位置などは、完成画像だけでは捉えにくい。WriteMirrorは、Surface Penから取得できる座標・時刻・ペン先圧等を用いて筆跡を再生し、本人の主観的な振り返りと、端末が直接観測した中立的な情報を並べるアプリである。平仮名・片仮名・漢字を対象としたWindowsの日本語手書き候補、音声案内、二回の試行比較を備える。",
)
add_para(
    doc,
    "先行研究は、オンライン筆跡が完成画像にない時間的・運動的情報を持つこと、児童のデジタル書字では成人データや紙筆条件をそのまま適用できないこと、表面に加わるペン圧と握筆力が同一ではないことを示している。本研究では、これらを踏まえ、診断や能力推定を避けた「書字プロセスの可視化」として技術を位置づける。成果は動作するARM64アプリ、研究計画書、テスト結果、デモ資料であり、教育的有効性の実証は範囲外とする。",
)

add_heading(doc, "目次", 1)
add_bullets(doc, [
    "1. 研究背景",
    "2. 先行研究と本研究への示唆",
    "3. 現状の問題と研究課題",
    "4. 研究目的・研究質問・設計仮説",
    "5. 提案アプリの概要",
    "6. 対象者と利用場面",
    "7. 問題解決の方法",
    "8. システム構成と実装",
    "9. データ利用とモデルの位置づけ",
    "10. 研究・評価方法",
    "11. 倫理・プライバシー・安全性",
    "12. 期待される効果",
    "13. 実施体制・工程・成果物",
    "14. リスク、限界、今後の研究可能性",
    "15. 結論",
    "参考文献",
])
doc.add_page_break()

add_heading(doc, "1. 研究背景", 1)
add_heading(doc, "1.1 書字学習における評価の偏り", 2)
add_para(
    doc,
    "学校や家庭の書字練習では、完成した文字が正しいか、形が整っているか、読めるかが中心になりやすい。この方法は成果を確認するには有効である一方、学習者がどこで迷ったか、どの筆画の前後で間があいたか、どの部分を本人が書きにくいと感じたかを扱いにくい。特に低学年や書字に困り感を持つ児童にとって、結果だけを訂正されることは、自分の書き方を理解する手掛かりにならない可能性がある。",
)
add_heading(doc, "1.2 デジタルペンが可能にする観察", 2)
add_para(
    doc,
    "電子ペンを用いると、完成画像に加え、座標、時刻、筆画の区切り、ペン先圧、傾き等を取得できる。これにより、書字を時間順に再生し、書き始めから書き終わりまでの過程を本人が確認できる。MicrosoftのWindows Inkは、位置と動きに加えてデジタイザーが取得するペン圧を扱えるため、Surfaceは書字プロセス研究の試作環境として利用できる。",
)
add_heading(doc, "1.3 本研究の社会的・教育的意義", 2)
add_para(
    doc,
    "本研究の意義は、児童を自動的に評価するAIを作ることではなく、本人が自分の書字過程を振り返るための情報設計を検討する点にある。主観的な感覚を先に尋ね、次に端末が観測した事実を示し、再試行につなげることで、「できた／できない」だけではない対話的な練習方法を提案する。また、保存しない既定モードにより、児童データを必要以上に残さない設計の実例を示す。",
)

add_heading(doc, "2. 先行研究と本研究への示唆", 1)
add_heading(doc, "2.1 オンライン・オフライン日本語筆跡", 2)
add_para(
    doc,
    "Seki（2020）は、401名から英字、数字、平仮名、片仮名、漢字を含む702文字を収集し、各文字を5回書くオンライン・オフライン併用データを構築した。オンラインデータには座標、時刻、ペン圧が含まれる。この研究は法科学的な筆者分類が目的であり、児童の教育支援に直接転用できるものではないが、同一人物による複数回試行と、軌跡・画像の二つの表現を併用する設計の有用性を示唆する。",
)
add_heading(doc, "2.2 児童筆跡における動的情報", 2)
add_para(
    doc,
    "Corbilléら（2020）は、児童のオンライン筆跡について、最終形状だけでなく、筆画の方向や角度を複数の画像チャネルとしてCNNへ入力すると認識性能が改善することを示した。児童の文字は成人より変形が大きく、成人データだけで学習した認識器では扱いにくい。WriteMirrorにとっては、広い文字認識と書字プロセス観察を分離しつつ、将来モデルを検討する場合は静止画像と時系列を併用すべきことを示す。",
)
add_heading(doc, "2.3 日本語児童のオンライン漢字評価", 2)
add_para(
    doc,
    "Inoueら（2024）は、Surface Go 2と電子ペンを用いた自己実施型の漢字評価OAHaSを開発し、6～12歳の日本語児童261名で紙筆テストとの関連を検証した。120漢字の学習用データには児童・青年・成人の筆跡と誤字例が含まれ、CNNによる自動判定は人工判定と高い一致を示した。一方、誤った文字を正しいと判定しない能力には限界があり、モデルが常に確定判断を返すべきではないことも読み取れる。また、研究実施時には保護者同意、児童のアセント、訓練を受けた補助者が用意されており、自己実施型であることと完全な無監督利用が同義ではない。",
)
add_heading(doc, "2.4 端末表面とペン圧の解釈", 2)
add_para(
    doc,
    "Alamargotら（2015）は、タブレット表面が低学年児童の筆画間の間や運動に影響し、その影響が学年によって異なることを報告した。また、Hochhauserら（2021/2023）の小規模研究では、握筆時の指の力とタブレット表面に加わるペン圧が異なる挙動を示した。したがって、Surfaceが返す0～1のペン圧値を、握る力、緊張、疲労、困難の直接指標として扱うことはできない。",
)
add_heading(doc, "2.5 筆跡からの過剰推定とプライバシー", 2)
add_para(
    doc,
    "近年の研究では、筆跡の速度、停止、修正等を教育的推論へ利用する提案がある一方、日本の児童・生徒の筆跡特徴から学年、性別、学業成績に関する信号を予測できる可能性も報告されている。これは筆跡が単なる文字画像ではなく、個人に関する情報を含み得ることを意味する。WriteMirrorは、能力、認知の深さ、性格、性別、学業成績、障害を推定しない方針を採用する。",
)

add_table(doc, ["先行研究から得た知見", "WriteMirrorへの反映", "採用しない解釈"], [
    ["同一人物の複数回試行", "2回比較と、任意の5回個人基線という設計案", "他者平均との差による順位づけ"],
    ["静止画像とオンライン軌跡の相補性", "画像認識と軌跡観察を別モジュールとして扱う", "単一スコアによる総合能力判定"],
    ["児童データの必要性", "成人データの結果を児童へ一般化しない", "成人モデルを児童向けAIと表示すること"],
    ["端末表面の影響", "Surface上の本人内比較に限定する", "紙筆常模との直接比較"],
    ["表面ペン圧と握筆力の差", "「端末が取得したペン先圧」と表現する", "握る力、緊張、疲労の推定"],
    ["筆跡の属性漏えい", "保存最小化と推論禁止項目を明文化する", "性別、学年、成績、診断の予測"],
], widths=[5.3, 6.7, 5.3])

add_heading(doc, "3. 現状の問題と研究課題", 1)
add_table(doc, ["現状の問題", "影響", "本研究の対応"], [
    ["完成文字中心の評価", "迷い、停止、書き直し等の過程が見えにくい", "軌跡再生と時間順の観察"],
    ["本人の感覚が残らない", "外部からの採点だけになりやすい", "書きやすさと気になった位置を本人が選択"],
    ["数値の過剰解釈", "遅さや圧力を能力・障害と誤認する危険", "観測事実と禁止推論を分離"],
    ["日本語文字種の広さ", "数文字だけの専用分類器では実用性が低い", "Windows日本語手書き候補と自由入力を併用"],
    ["児童データの機微性", "保存・共有・再学習でプライバシーリスクが生じる", "独立練習は保存しない。研究収集は実施しない"],
    ["低学年向け操作の複雑さ", "説明が読めない、誤操作から戻れない可能性", "短い日本語、音声、段階表示、大きい操作領域"],
])
add_para(
    doc,
    "以上から、本研究の中心課題を「児童を採点・診断せず、本人の主観とSurface Penから得られる観測事実を結び付け、書字プロセスの振り返りと再試行を支援する技術的仕組みを構築できるか」と定義する。",
)

add_heading(doc, "4. 研究目的・研究質問・設計仮説", 1)
add_heading(doc, "4.1 研究目的", 2)
add_para(
    doc,
    "本研究の目的は、Surface Penを用いて書字過程を記録・可視化し、本人の主観的振り返りと中立的な観測情報を安全に提示する日本語アプリの技術的実現可能性を示すことである。併せて、公開データ、手書き認識、AI分析を教育支援へ利用する際の適用範囲と倫理的境界を整理する。",
)
add_heading(doc, "4.2 研究質問", 2)
add_bullets(doc, [
    "RQ1：Surface Penの筆跡を、本人が理解できる形で再生・可視化できるか。",
    "RQ2：本人の「書きにくかった場所」と端末上の時間的イベントを、採点せずに並べて提示できるか。",
    "RQ3：平仮名・片仮名・漢字について、手動入力とWindows文字候補を組み合わせた柔軟な操作を構成できるか。",
    "RQ4：既定で保存しない設計、音声案内、短い日本語により、児童独立利用を想定した候補UIを構成できるか。",
    "RQ5：先行研究の知見を取り入れながら、診断・能力推定・因果推論を避けた説明方針を定義できるか。",
])
add_heading(doc, "4.3 設計仮説", 2)
add_bullets(doc, [
    "H1：完成画像だけでなく軌跡を再生すると、書いた順序や間を本人が確認しやすくなる。",
    "H2：数値より先に本人の感覚を尋ねると、システムの値に回答を誘導されにくくなる。",
    "H3：観測事実、解釈の限界、次に試す問いを分けると、採点的なフィードバックを避けやすい。",
    "H4：独立練習でデータを保存しなければ、授業内デモにおける不要な個人データ保持を減らせる。",
])
add_note(doc, "仮説の扱い", "本授業課題では実児童を対象とした比較実験を行わないため、H1～H4は実証仮説ではなく、プロトタイプ設計を導く作業仮説として扱う。", fill=PALE_GREEN)

add_heading(doc, "5. 提案アプリの概要", 1)
add_heading(doc, "5.1 コンセプト", 2)
add_para(
    doc,
    "WriteMirrorは、書字の「鏡」として、書いた結果を評価するのではなく、書いていた過程を本人へ返す。中核は、①文字を選ぶ、②書く、③感覚を選ぶ、④軌跡を見る、⑤もう一度試す、という短い循環である。",
)
add_flow(doc)
add_heading(doc, "5.2 現在実装されている主要機能", 2)
add_table(doc, ["機能", "内容", "研究上の意味"], [
    ["Surface Pen入力", "座標、時刻、ペン先圧、傾き、筆画を取得", "完成画像と書字過程を分けて扱う"],
    ["軌跡再生", "書いた順序を時間に沿って再表示", "本人が過程を振り返る"],
    ["主観回答", "気持ちを選び、必要な場合だけ位置を囲む", "肯定・問題なし・回答拒否を観測候補で上書きしない"],
    ["中立的観察", "総時間、筆画間空白、本人が囲んだ位置との関係", "診断せずに観測事実を示す"],
    ["二回比較", "本人の一回目と二回目のみを比較", "他者順位ではなく自己内の変化を見る"],
    ["日本語文字候補", "Windowsの平仮名・片仮名・漢字候補と自由入力", "三文字だけに限定しない"],
    ["音声案内", "説明と結果を日本語音声で読み上げ", "低学年、読字負担への配慮"],
    ["保存方針", "ひとりで練習は保存なし。共同確認は明示選択時のみ保存", "データ最小化"],
])

add_heading(doc, "6. 対象者と利用場面", 1)
add_heading(doc, "6.1 想定する対象者", 2)
add_table(doc, ["区分", "対象", "想定するニーズ"], [
    ["主対象", "小学校低学年、書字に困り感を持つ可能性のある児童", "結果だけでなく、自分の書き方を見て確かめたい"],
    ["一般対象", "年齢や診断の有無を問わない日本語学習者", "平仮名・片仮名・漢字を自由に練習したい"],
    ["補助利用者", "教員、保護者、支援者", "本人の感覚を起点に会話する材料がほしい"],
    ["本研究の評価者", "小グループメンバー、教員、成人デモ参加者", "技術的完成度、説明の一貫性、操作性を確認する"],
])
add_heading(doc, "6.2 利用場面", 2)
add_bullets(doc, [
    "授業発表・展示：Surface実機上でペン入力、再生、認識候補、音声を短時間で示す。",
    "個人練習のデモ：保存しないモードで、本人が一人で二回試す流れを体験する。",
    "教員・保護者への説明：数値を診断に使わず、本人の感覚を聞くための材料として示す。",
    "研究・技術教育：デジタルインク、ローカル処理、データ最小化、AIの限界を学ぶ題材とする。",
])
add_heading(doc, "6.3 本研究で扱わない場面", 2)
add_bullets(doc, [
    "実児童を対象とした臨床評価、障害スクリーニング、治療・訓練効果の判定",
    "成績評価、学校内順位づけ、児童間比較、教師の評価を代替する自動採点",
    "家庭・学校への常設配布、大規模データ収集、クラウドへの筆跡送信",
    "Phi Silica、NPU、独自AIモデルが実接続済みであるという主張",
])

add_heading(doc, "7. 問題解決の方法", 1)
add_heading(doc, "7.1 主観を先に扱う", 2)
add_para(
    doc,
    "システムが数値や候補を示す前に、利用者自身が「ためらった」「書きにくかった」「気になった」「特になし」「うまくいった」「答えない」から選び、必要な場合だけ気になった位置を囲む。「特になし」「うまくいった」「答えない」では候補を自動表示せず、本人が「観測を見てみる」または「これで終わる」を選ぶ。これにより、分析結果が本人の回答を上書きする構造を避ける。",
)
add_heading(doc, "7.2 観測事実と解釈を分離する", 2)
add_para(
    doc,
    "表示可能な情報は、座標、筆画数、筆画間の空白時間、端末が取得したペン先圧等に限定する。WPFは筆画首尾の時刻を点列へ等間隔補間するため、局所速度、低速度、最遅区間を表示・照合・音声・フィードバックへ使用しない。画間は未観測の空中経路を線で結ばず、前画終点と次画始点だけを示す。表示文には「困難、原因、能力、改善を意味しない」ことを添える。",
)
add_heading(doc, "7.3 文字認識と書字分析を分離する", 2)
add_para(
    doc,
    "文字認識は「何という文字に見えるか」の候補を返す機能であり、書字分析は「どのような時間順で書いたか」を観察する機能である。Windowsの認識候補は便利な補助だが、正誤判定や教育的評価ではない。候補が違う場合、利用者は自由に修正できる。",
)
add_heading(doc, "7.4 本人内比較に限定する", 2)
add_para(
    doc,
    "現在の比較は同一端末・同一人物・同一セッションの二回に限定する。二回目が速い、圧力変動が小さい等の差があっても、練習効果や改善とは呼ばない。慣れ、順序、速度と正確さのトレードオフ等、別の説明があり得るためである。",
)

add_heading(doc, "8. システム構成と実装", 1)
add_heading(doc, "8.1 構成", 2)
add_table(doc, ["層", "主要要素", "役割"], [
    ["UI", "WriteMirror.Wpf", "日本語画面、ペン入力、音声、段階フロー"],
    ["Core", "入力記録、分析、比較、安全なフィードバック", "UIから独立した計測・方針"],
    ["Recognizer", "Windows日本語手書き認識ヘルパー", "平仮名・片仮名・漢字の候補生成"],
    ["Infrastructure", "JSONセッション保存", "明示同意された共同確認時だけ保存"],
    ["Experiment", "KanjiVG軌跡実験", "独立した研究用試作。正式アプリには未統合"],
])
add_heading(doc, "8.2 対象環境", 2)
add_bullets(doc, [
    "端末：Microsoft Surface Pro 11、Snapdragon X Elite、ARM64",
    "入力：Surface Slim Pen 2。右手・左手を本人が選択可能",
    "安定版UI：WPF / Windows Ink / ARM64",
    "表示・音声：日本語（ja-JP）、Windows日本語音声",
    "動作方針：オフラインで主要機能を完結し、外部AIサービスを必須としない",
])
add_heading(doc, "8.3 現在の技術的到達点", 2)
add_table(doc, ["項目", "確認済みの状態", "主張しないこと"], [
    ["Coreテスト", "73件すべて成功", "教育的効果や臨床妥当性"],
    ["WPFビルド", "ARM64 / x64で警告0、エラー0", "全Surface機種での互換性"],
    ["実機デモ", "Surface Pen入力、再生、文字候補、音声が利用可能", "児童による無監督利用の実証"],
    ["軌跡学習実験", "KanjiVGでノイズ除去型GRUの実行を確認", "文字認識・障害判定・NPU実行"],
])
add_heading(doc, "8.4 Vibe Codingの利用", 2)
add_para(
    doc,
    "本開発では、生成AIをコード案、テスト案、文書草案、レビュー観点の作成に利用した。ただし、生成物はそのまま正しいとはみなさず、ビルド、単体テスト、実機確認、文献照合によって検証した。開発時にAIを使うことと、完成アプリが実行時AIを搭載していることは区別する。",
)

add_heading(doc, "9. データ利用とモデルの位置づけ", 1)
add_heading(doc, "9.1 公開データの候補と制約", 2)
add_table(doc, ["資料・データ", "利用可能な内容", "本研究での扱い"], [
    ["KanjiVG", "日本の漢字の筆順・方向を含むSVG。CC BY-SA 3.0", "参照軌跡と学習パイプラインの動作確認に使用"],
    ["ETL Character Database", "平仮名、片仮名、教育漢字等のオフライン画像", "先行研究・将来候補。追加学習は実施しない"],
    ["HANDS-nakayosi", "163名、4438文字種の実測オンライン日本語筆跡", "利用申請・ライセンスが必要。取得しない"],
    ["Seki 2020", "401名、702文字、各5回、オンライン・オフライン", "論文上の設計参考。データは取得しない"],
    ["OAHaS", "児童を含む120漢字の画像・誤字例・評価資料", "公開状況を確認。再利用許諾が明確でないため取得・製品学習しない"],
])
add_heading(doc, "9.2 現在の軌跡モデル実験", 2)
add_para(
    doc,
    "現在の実験モデルは、KanjiVGの整形済み軌跡を入力し、座標へ小さなノイズを加えて元軌跡を再構成する双方向GRUである。入力特徴は正規化座標と筆画末尾であり、児童の時刻、ペン先圧、傾き、困難ラベルを含まない。このモデルはパイプラインの動作確認用であり、文字認識器でも児童分析器でもない。正式アプリには接続しない。",
)
add_heading(doc, "9.3 本研究でモデルを追加訓練しない理由", 2)
add_bullets(doc, [
    "小グループ作業の目的は研究試作と説明可能なデモであり、大規模モデル開発ではない。",
    "公開されていることと、製品学習・再配布が許諾されていることは同じではない。",
    "成人筆跡や理想筆順だけでは、低学年児童の多様な筆跡を代表できない。",
    "教育的・臨床的ラベルのないデータから、困難や改善を学習することはできない。",
    "不十分なモデルを接続すると、Windows既存認識より精度・安全性が下がる可能性がある。",
])
add_heading(doc, "9.4 実行時AIを追加しない判断", 2)
add_para(
    doc,
    "Microsoftは、Phi Silicaを新しいオンデバイスモデルAion Instructへ置き換える移行計画を公表している。本課題は8月1日の固定デモを目的とし、実行時AIの有用性・安全性・NPU実行証拠を検証する時間と必要性がない。そのため、Phi Silicaへの駆け込み実装もAion Instructの先行統合も行わず、実装済みのWindows日本語手書き文字候補と固定日本語テンプレートだけを正確に説明する。",
)

add_heading(doc, "10. 研究・評価方法", 1)
add_heading(doc, "10.1 研究デザイン", 2)
add_para(
    doc,
    "本研究は、先行研究調査、要求定義、反復的プロトタイプ開発、技術テスト、成人による実機確認を組み合わせたデザイン研究である。統制群を用いた教育効果実験や、児童を対象とする人参加研究ではない。",
)
add_heading(doc, "10.2 評価対象", 2)
add_bullets(doc, [
    "機能完全性：起動、入力、再生、主観回答、比較、認識候補、音声、削除",
    "技術安定性：ARM64ビルド、単体テスト、連続操作、例外処理",
    "表現の安全性：診断語、能力語、因果表現、未取得の数値を出さないこと",
    "プライバシー：独立練習終了後に復元可能なセッションファイルが残らないこと",
    "アクセシビリティ候補：キーボード操作、100%表示、大きな操作領域、日本語音声",
    "文書整合性：アプリ、README、研究計画書、発表資料で主張が一致すること",
])
add_heading(doc, "10.3 技術的受入基準", 2)
add_table(doc, ["評価項目", "受入基準", "証拠"], [
    ["ビルド", "WPF ARM64 Releaseで警告0・エラー0", "ビルドログ"],
    ["Core", "単体テスト73/73成功", "テストログ"],
    ["ペン入力", "Surface Penで連続した筆跡を記録し、タッチ誤入力を区別", "実機確認"],
    ["日本語", "主要画面、説明、フィードバックが日本語", "画面確認"],
    ["文字種", "平仮名・片仮名・漢字を自由入力・修正可能", "代表文字による確認"],
    ["音声", "説明と結果を日本語で読み上げ可能", "実機確認"],
    ["保存", "独立練習ではセッションファイルを新規作成しない", "保存先確認・単体テスト"],
    ["安全表現", "障害・診断・能力・正常／異常を出力しない", "テスト・文言レビュー"],
])
add_heading(doc, "10.4 収集しない評価", 2)
add_para(
    doc,
    "本研究では、児童の成績、診断、性別、年齢、学校名、個人識別情報、長期筆跡履歴を収集しない。また、アプリによって書字能力が改善したか、困難が軽減したかを測定しない。したがって、発表では「改善した」「障害を発見できる」「誰でも一人で安全に使える」と結論づけない。",
)

add_heading(doc, "11. 倫理・プライバシー・安全性", 1)
add_heading(doc, "11.1 児童を想定した研究倫理", 2)
add_para(
    doc,
    "児童の筆跡は個人差を含む行動データであり、研究利用には目的、保存、利用者、削除、第三者提供を明確にする必要がある。本研究は実児童を対象にデータ収集しないため、人参加研究の同意手続きを実施しない。将来、児童評価へ進む場合には、所属機関の倫理手続き、保護者の説明と同意、児童本人が理解できるアセント、途中中止、データ削除を別途必要とする。",
)
add_heading(doc, "11.2 データ最小化", 2)
add_bullets(doc, [
    "「ひとりで練習」を既定とし、筆跡をプロセス外へ保存しない。",
    "「いっしょに確認」でも、保存はセッションごとの明示選択とする。",
    "クラウド送信、広告、アカウント、児童プロファイル、継続追跡を行わない。",
    "削除操作は現在のセッションだけを明確に対象とし、復元可能性を誤って説明しない。",
])
add_heading(doc, "11.3 推論禁止項目", 2)
add_table(doc, ["扱ってよい表現", "条件付き表現", "禁止する推論"], [
    ["第2画と第3画の間が他の間より長かった", "端末が取得したペン先圧が変化した", "握る力が弱い／強い"],
    ["本人がこの位置を気になった場所として囲んだ", "二回の総時間に差があった", "能力が改善した／低下した"],
    ["Windowsの文字候補は『木』だった", "候補が一致しなかった", "書字障害、発達障害、異常"],
    ["この端末上の本人内観察", "次に試す問い", "性格、緊張、努力、学年、性別、成績"],
])
add_heading(doc, "11.4 心理的安全", 2)
add_bullets(doc, [
    "赤い警告、失敗音、点数、ランキング、正常／異常表示を使わない。",
    "候補が違っても利用者が修正でき、失敗扱いにしない。",
    "いつでも中止・削除できる操作を用意する。",
    "結果は短い観察、本人の振り返り、次に試す問いの三部構成とする。",
    "肯定・問題なし・回答拒否では客観候補を自動表示せず、通常終了を再試行の失敗として扱わない。",
])

add_heading(doc, "12. 期待される効果", 1)
add_heading(doc, "12.1 利用者に期待される効果", 2)
add_bullets(doc, [
    "完成した文字だけでなく、自分が書いた順序や間を視覚的に確認できる。",
    "「どこが気になったか」を本人の言葉や位置で表現するきっかけになる。",
    "他者との比較や点数ではなく、自分の二回の試行を落ち着いて見比べられる。",
    "音声案内と短い日本語によって、文章を読む負担を減らせる可能性がある。",
])
add_heading(doc, "12.2 教員・保護者・支援者に期待される効果", 2)
add_bullets(doc, [
    "「正しい／間違い」だけでなく、本人がどこを気にしたかを尋ねる材料になる。",
    "完成画像からは分からない時間順の情報を、会話の補助として確認できる。",
    "数値を診断へ結び付けないデータリテラシーの例を示せる。",
])
add_heading(doc, "12.3 小グループ研究として期待される成果", 2)
add_bullets(doc, [
    "Surface Pen、Windows Ink、音声合成、ARM64を統合した動作試作を提示できる。",
    "公開データとAIモデルの利用可能性・ライセンス・限界を説明できる。",
    "児童向けアプリで求められるデータ最小化と非診断設計を具体化できる。",
    "Vibe Codingで生成したコードを、テストと文献によって検証する実践例になる。",
])
add_note(doc, "効果の限界", "上記は期待される効果であり、本研究では児童参加実験や教育成果測定を行わないため、効果量、改善率、診断精度として報告しない。", fill=PALE_GREEN)

add_heading(doc, "13. 実施体制・工程・成果物", 1)
add_heading(doc, "13.1 小グループの役割", 2)
add_table(doc, ["役割", "主な担当", "レビュー責任"], [
    ["企画・先行研究", "研究背景、文献、課題、倫理境界", "主張と出典の一致"],
    ["アプリ実装", "UI、Surface Pen、認識候補、音声", "実機動作と例外処理"],
    ["Core・テスト", "記録、分析、比較、保存方針", "再現可能な単体テスト"],
    ["文書・発表", "研究計画書、README、デモ台本", "日本語と表現の一貫性"],
])
add_heading(doc, "13.2 工程", 2)
add_table(doc, ["段階", "内容", "状態"], [
    ["1. 課題設定", "書字結果ではなくプロセスを扱う研究課題の定義", "完了"],
    ["2. 先行研究", "児童筆跡、日本語データ、端末表面、圧力、プライバシーの調査", "完了"],
    ["3. 試作", "WPF ARM64アプリ、Surface Pen、音声、文字候補、保存方針", "完了"],
    ["4. 安全修正", "診断・改善表現の削除、独立練習の保存禁止、児童向け段階UI", "完了"],
    ["5. 技術検証", "Coreテスト、ARM64ビルド、成人による実機デモ", "完了／発表時確認"],
    ["6. 報告", "研究計画書、成果説明、90秒デモ", "最終段階"],
])
add_heading(doc, "13.3 成果物", 2)
add_bullets(doc, [
    "WriteMirror 0.5.2 ARM64実行版",
    "日本語アプリ開発・研究計画書（本書）",
    "Core単体テストおよびビルド結果",
    "先行研究と公開データの利用条件整理",
    "実機デモ手順および安全な説明文",
])

add_heading(doc, "14. リスク、限界、今後の研究可能性", 1)
add_heading(doc, "14.1 主なリスクと対応", 2)
add_table(doc, ["リスク", "本研究への影響", "対応"], [
    ["観測値の過剰解釈", "児童の困難や改善と誤認される", "非診断表示、禁止語、安全バリデーション"],
    ["Windows認識候補の誤り", "違う文字が自動入力される", "候補は補助と明記し、自由修正を維持"],
    ["タブレット固有の影響", "紙筆や他端末へ一般化できない", "同一Surface上の本人内観察に限定"],
    ["公開データの権利", "再学習・配布が不適切になる", "利用規約を個別確認し、本研究では追加取得しない"],
    ["児童UIの未検証", "実際の理解・誤操作が不明", "児童独立利用を実証済みと表現しない"],
    ["AI/NPUの誤表示", "技術的主張が事実と不一致になる", "未接続と明記し、デモ対象外とする"],
])
add_heading(doc, "14.2 本研究の限界", 2)
add_bullets(doc, [
    "児童を対象とした可用性、心理的安全性、教育効果を測定していない。",
    "Windowsの手書き認識は候補を返すが、すべての手書き日本語を正確に認識する保証はない。",
    "二回の試行差は信頼性のある個人基線ではなく、改善の証拠ではない。",
    "ペン先圧はデジタイザーの相対値であり、握筆力や身体負担を表さない。",
    "KanjiVGモデルは理想軌跡上の去雑音実験であり、児童筆跡への性能を示さない。",
])
add_heading(doc, "14.3 今後の研究可能性（本課題では実施しない）", 2)
add_bullets(doc, [
    "児童・保護者・教員を含む倫理承認済みの形成的可用性調査",
    "同一人物による5回程度の個人基線と不確実性表示",
    "児童筆跡を含む許諾済みデータによる画像・時系列の二系統モデル",
    "書き手、学校、端末を分離した評価と、属性漏えい監査",
    "音声、キーボード、拡大、高コントラストを含むアクセシビリティ評価",
])
add_para(
    doc,
    "これらは研究を継続する場合の可能性を示すものであり、本小グループ作業の未完了タスクではない。現在の成果は、限定された範囲を正確に説明できる技術試作として完結させる。",
)

add_heading(doc, "15. 結論", 1)
add_para(
    doc,
    "WriteMirrorは、Surface Penから得られる書字過程を再生し、本人の感覚と中立的な観測事実を結び付ける日本語研究試作である。先行研究は、オンライン軌跡の価値を支持する一方、児童データ、端末表面、ペン圧、属性漏えい、誤判定に慎重な扱いが必要であることも示す。本研究はその両面を取り入れ、診断・採点・能力推定を行わず、保存を最小化し、平仮名・片仮名・漢字、音声、右手・左手に配慮した技術構成を提案した。",
)
add_para(
    doc,
    "小グループ作業としての到達点は、教育効果を証明することではなく、問題設定、先行研究、設計、実装、テスト、倫理境界を一つの説明可能な成果へまとめたことである。発表では、動作するアプリと同時に「何ができ、何はまだ言えないか」を明確に示す。",
)

doc.add_page_break()
add_heading(doc, "参考文献", 1)
add_reference(doc, 1, "Seki, Y. (2020). Collection of Online and Offline Handwritten Japanese Characters and Handwriting Classification Using the Data. ICFHR 2020.", "https://doi.org/10.1109/ICFHR2020.2020.00056")
add_reference(doc, 2, "Corbillé, S., Fromont, É., Anquetil, É., & Nerdeux, P. (2020). Integrating Writing Dynamics in CNN for Online Children Handwriting Recognition. ICFHR 2020.", "https://doi.org/10.1109/ICFHR2020.2020.00057")
add_reference(doc, 3, "Inoue, T., Chen, Y., & Ohyanagi, T. (2024). Assessing handwriting skills in a web browser: Development and validation of an automated online test in Japanese Kanji. Behavior Research Methods.", "https://doi.org/10.3758/s13428-024-02562-6")
add_reference(doc, 4, "Alamargot, D., Morin, M.-F., & Simard-Dupuis, É. (2015). Does handwriting on a tablet screen affect students’ graphomotor execution? Human Movement Science, 44, 32–41.", "https://doi.org/10.1016/j.humov.2015.08.011")
add_reference(doc, 5, "Hochhauser, M., Wagner, M., & Shvalb, N. (2023). Assessment of children's writing features: A pilot method study of pen-grip kinetics and writing surface pressure. Assistive Technology, 35(1), 107–115.", "https://doi.org/10.1080/10400435.2021.1956640")
add_reference(doc, 6, "Nakamoto, R., Flanagan, B., Nakamura, K., & Ogata, H. (2026). Explain-from-Stroke: Capturing Invisible Learning Processes Through Handwriting Dynamics Analysis. AAAI 2026.", "https://doi.org/10.1609/aaai.v40i48.42118")
add_reference(doc, 7, "Iste, A., et al. (2026). Prediction of Grade, Gender, and Academic Performance of Children and Teenagers from Handwriting Using the Sigma-Lognormal Model. arXiv.", "https://arxiv.org/abs/2603.11519")
add_reference(doc, 8, "National Institute of Advanced Industrial Science and Technology. ETL Character Database.", "https://etlcdb.db.aist.go.jp/the-etl-character-database/")
add_reference(doc, 9, "Tokyo University of Agriculture and Technology, Nakagawa Laboratory. HANDS-nakayosi_t-98-09.", "https://web.tuat.ac.jp/~nakagawa/database/en/about_nakayosi.html")
add_reference(doc, 10, "KanjiVG Project. Kanji Vector Graphics. CC BY-SA 3.0.", "https://kanjivg.tagaini.net/")
add_reference(doc, 11, "Microsoft. Pen interactions and Windows Ink in Windows apps.", "https://learn.microsoft.com/en-us/windows/uwp/ui-input/pen-and-stylus-interactions")
add_reference(doc, 12, "Microsoft. Accessibility overview for Windows apps.", "https://learn.microsoft.com/en-us/windows/apps/design/accessibility/accessibility-overview")
add_reference(doc, 13, "文部科学省. 教育データの利活用に係る留意事項.", "https://www.mext.go.jp/a_menu/other/data_00007.htm")
add_reference(doc, 14, "個人情報保護委員会. 個人情報の保護に関する法律についてのガイドライン（通則編）.", "https://www.ppc.go.jp/personalinfo/legal/guidelines_tsusoku/")
add_reference(doc, 15, "Microsoft. Get started with Phi Silica in the Windows App SDK: transition guidance to Aion Instruct.", "https://learn.microsoft.com/en-us/windows/ai/apis/phi-silica")

add_heading(doc, "付録A：発表時の主張範囲", 1)
add_table(doc, ["発表してよいこと", "条件付きで説明すること", "発表しないこと"], [
    ["Surface Penで書字過程を記録・再生できる", "Windowsが日本語文字候補を返す", "すべての文字を100%認識できる"],
    ["本人の主観回答と観測位置を並べられる", "二回の値に差が観測された", "書字能力が改善した"],
    ["独立練習は保存しない設計である", "児童独立利用を想定した候補UIである", "児童だけで安全に使えることを実証した"],
    ["KanjiVGで学習実験を実行した", "将来は別方式のモデルを検討できる", "独自AI・Phi Silica・NPUがアプリで動いている"],
])

add_heading(doc, "付録B：90秒実機デモ", 1)
add_bullets(doc, [
    "0～15秒：研究課題を説明する。「完成文字の採点ではなく、書いている途中を振り返るアプリです」。",
    "15～35秒：任意の平仮名・片仮名・漢字を選び、Surface Penで書く。",
    "35～50秒：書きやすさを選び、必要なら気になった位置を囲む。",
    "50～65秒：軌跡再生と中立的な観察を示す。数値は困難や能力を意味しないと説明する。",
    "65～78秒：Windows日本語文字候補と音声読み上げを示す。候補は修正できる。",
    "78～90秒：独立練習では保存しないこと、児童効果は未検証であることを説明して終える。",
])

doc.save(OUTPUT)
print(OUTPUT)
